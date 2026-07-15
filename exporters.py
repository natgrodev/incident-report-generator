"""
Export a Markdown report to DOCX and PDF.

The model returns Markdown. Stakeholders want a document. This bridges the two.

Deliberately hand-rolled rather than using a Markdown-to-DOCX library: incident
reports use a narrow, predictable subset of Markdown (headings, tables, prose,
the occasional blockquote), and controlling the output formatting directly is
simpler than fighting a general-purpose converter.
"""

import io
import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
)


# ---------------------------------------------------------------- parsing

def _parse(markdown):
    """
    Turn the model's Markdown into a flat list of blocks.

    Each block is a tuple:
        ("h1" | "h2", text)
        ("p", text)
        ("quote", text)
        ("table", [[cell, cell], [cell, cell], ...])   # first row = header
    """
    blocks = []
    lines = markdown.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        # Table: a line starting with | begins one
        if line.lstrip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                raw = lines[i].strip().strip("|")
                cells = [c.strip() for c in raw.split("|")]
                # skip the |---|---| separator row
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
                    # clean each cell: the model sometimes wraps paths and tags in
                    # backticks, which render literally in DOCX and PDF
                    rows.append([_clean(c) for c in cells])
                i += 1
            if rows:
                blocks.append(("table", rows))
            continue

        if line.startswith("## "):
            blocks.append(("h2", _clean(line[3:])))
        elif line.startswith("# "):
            blocks.append(("h1", _clean(line[2:])))
        elif line.startswith("> "):
            blocks.append(("quote", _clean(line[2:])))
        elif line.startswith("---"):
            pass  # horizontal rules: skip
        else:
            blocks.append(("p", _clean(line)))

        i += 1

    return blocks


def _clean(text):
    """Strip Markdown emphasis markers — we apply formatting structurally instead."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text.strip()


# ---------------------------------------------------------------- DOCX

def to_docx(markdown, title="Post-Incident Report"):
    """Render the report as a .docx file. Returns bytes."""
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for kind, content in _parse(markdown):

        if kind in ("h1", "h2"):
            doc.add_heading(content, level=1 if kind == "h1" else 2)

        elif kind == "p":
            doc.add_paragraph(content)

        elif kind == "quote":
            p = doc.add_paragraph(content)
            p.paragraph_format.left_indent = Pt(18)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        elif kind == "table":
            rows = content
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=0, cols=cols)
            table.style = "Light Grid Accent 1"

            # A "key/value" table (like Incident Details) has no header row — its
            # first row is real data. Detect it so we don't bold and shade a data row.
            is_key_value = bool(rows) and rows[0] and rows[0][0].strip().lower() == "incident title"

            for idx, row in enumerate(rows):
                cells = table.add_row().cells
                for c_idx in range(cols):
                    text = row[c_idx] if c_idx < len(row) else ""
                    cells[c_idx].text = text
                    # Bold the first column of a key/value table (the field names),
                    # or the whole first row of a normal table (the header).
                    bold = (c_idx == 0) if is_key_value else (idx == 0)
                    if bold:
                        for para in cells[c_idx].paragraphs:
                            for run in para.runs:
                                run.bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------- PDF

def to_pdf(markdown, title="Post-Incident Report"):
    """Render the report as a PDF. Returns bytes."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
        title=title,
    )

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "body", parent=styles["Normal"],
        fontSize=9.5, leading=13, spaceAfter=6,
    )
    h1 = ParagraphStyle(
        "h1", parent=styles["Heading1"],
        fontSize=16, spaceBefore=12, spaceAfter=8,
    )
    h2 = ParagraphStyle(
        "h2", parent=styles["Heading2"],
        fontSize=12, spaceBefore=10, spaceAfter=6,
    )
    quote = ParagraphStyle(
        "quote", parent=body,
        leftIndent=14, textColor=colors.HexColor("#555555"),
        fontName="Helvetica-Oblique",
    )

    story = [Paragraph(title, styles["Title"]), Spacer(1, 10)]

    for kind, content in _parse(markdown):

        if kind == "h1":
            story.append(Paragraph(content, h1))
        elif kind == "h2":
            story.append(Paragraph(content, h2))
        elif kind == "p":
            story.append(Paragraph(content, body))
        elif kind == "quote":
            story.append(Paragraph(content, quote))
        elif kind == "table":
            rows = content
            cols = max(len(r) for r in rows)

            # Wrap every cell in a Paragraph so long text wraps instead of overflowing
            cell_style = ParagraphStyle("cell", parent=body, fontSize=8.5, leading=11)
            data = []
            for row in rows:
                padded = list(row) + [""] * (cols - len(row))
                data.append([Paragraph(c, cell_style) for c in padded])

            # Key/value table (Incident Details): no header row, shade the field column.
            is_key_value = bool(rows) and rows[0] and rows[0][0].strip().lower() == "incident title"

            avail = doc.width
            if is_key_value:
                # Narrow field column, wide value column
                col_widths = [avail * 0.32, avail * 0.68]
            else:
                col_widths = [avail / cols] * cols

            table = Table(data, colWidths=col_widths, repeatRows=0 if is_key_value else 1)

            base_style = [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#999999")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
            if is_key_value:
                # Shade the field-name column instead of a header row
                base_style.append(("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f0f0f0")))
            else:
                base_style.append(("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8e8e8")))

            table.setStyle(TableStyle(base_style))
            story.append(table)
            story.append(Spacer(1, 8))

    doc.build(story)
    return buffer.getvalue()
